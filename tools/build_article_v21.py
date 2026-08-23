"""Aplica as correções metodológicas da v2.1 sobre uma cópia da v2.

O script preserva a v2 e produz um DOCX de trabalho. Contagens, PRISMA,
figuras e tabelas só serão atualizados após o fechamento da matriz autoritativa.
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


SOURCES = (
    "A busca combinou consultas estruturadas e expansão bibliográfica. A recuperação ampliada utilizou "
    "OpenAlex, Crossref, Europe PMC, arXiv e DOAJ em cinco famílias: governança de LLMs, LLMOps e "
    "observabilidade, governança conversacional, ambientes regulados e supervisão humana/contestabilidade. "
    "Referências, citações, autoria e veículos foram rastreados para ampliar a cobertura."
)

RETRIEVAL = (
    "A recuperação bibliográfica foi realizada em múltiplas fontes e famílias de consulta. A análise de sensibilidade "
    "examinou faixas sucessivas até a centésima posição por combinação, complementadas por rastreamento de referências, "
    "citações, autoria e veículos. Os registros "
    "foram deduplicados por identificadores persistentes e, subsidiariamente, por título normalizado, autoria e "
    "ano. A abrangência foi examinada por análise de sensibilidade das posições recuperadas e pelo confronto entre "
    "fontes, mantendo-se no corpus apenas estudos elegíveis segundo os critérios definidos."
)

SUMMARY_PT = (
    "Esta revisão sistemática investiga mecanismos de governança conversacional em sistemas baseados em LLMs "
    "aplicados a ambientes regulados. Orientada pelo PRISMA 2020, documentou decisões sobre 383 textos completos e "
    "consolidou 358 estudos únicos no corpus analítico, sendo 30 evidências centrais e 328 de apoio. Os resultados "
    "mostram predominância de compliance, gestão de risco, accountability e monitoramento, enquanto contestabilidade "
    "e reparo permanecem pouco consolidados. A principal contribuição é um modelo integrado de cinco camadas - técnica, "
    "interacional, organizacional, regulatória e evolutiva - derivado da síntese e ainda dependente de validação empírica."
)

SUMMARY_EN = (
    "This systematic review investigates conversational governance mechanisms in LLM-based systems deployed in "
    "regulated environments. Guided by PRISMA 2020, it documented decisions for 383 full texts and consolidated 358 "
    "unique studies in the analytical corpus, including 30 central and 328 supporting evidence studies. Results show "
    "a predominance of compliance, risk management, accountability, and monitoring, whereas contestability and repair "
    "remain underdeveloped. The main contribution is an integrated five-layer model - technical, interactional, "
    "organizational, regulatory, and evolutionary - derived from the synthesis and still requiring empirical validation."
)

SELECTION = (
    "O snowballing incluiu referências citadas, trabalhos citantes, relacionados e expansões controladas por autoria "
    "e veículo. Todos os registros foram consolidados e deduplicados por DOI, título exato e similaridade textual, "
    "com validação de grupos potencialmente ambíguos. Como verificação de abrangência, a execução ampliada registrou "
    "1.342 ocorrências e 1.074 registros únicos após deduplicação interna, distribuídos por todas as faixas examinadas "
    "até a centésima posição. A seleção integral documentou decisões sobre 383 textos completos."
)

CORPUS = (
    "Os critérios de elegibilidade selecionaram estudos sobre LLMs, IA generativa ou sistemas conversacionais que "
    "apresentassem mecanismos de governança e implicação para ambientes regulados ou de alto impacto. Após a decisão "
    "integral e a reconciliação de versões, o corpus analítico reuniu 358 estudos únicos: 30 evidências centrais e "
    "328 evidências de apoio. Vinte e quatro registros permaneceram fora do escopo analítico e uma versão redundante "
    "foi vinculada à publicação correspondente, sem exclusão do histórico documental."
)

TRACEABILITY = (
    "Os 358 estudos estão identificados individualmente no suplemento por referência, classificação, setor, família "
    "de mecanismos, camada normalizada, PDF, hash, página e trecho de evidência. A matriz longa estudo-mecanismo-camada "
    "é a fonte única para recalcular as tabelas e figuras; todos os arquivos indicados foram localizados e nenhum estudo "
    "incluído permaneceu sem evidência verificável."
)

CENTRAL_CRITERIA = (
    "A classificação como evidência central exigiu simultaneamente três condições: tratamento direto de "
    "governança, supervisão, risco, accountability, auditoria, compliance ou operação controlada de LLMs ou "
    "sistemas conversacionais; relação explícita com ambiente regulado, de alto impacto ou mecanismo "
    "demonstravelmente transferível; e contribuição substantiva para pelo menos uma questão da revisão, por "
    "resultado empírico, síntese sistemática, mecanismo avaliado ou arquitetura conceitual. Estudos elegíveis "
    "com contribuição periférica, contextual ou apenas transferível foram classificados como evidência de apoio."
)

LAYER_ORIGIN = (
    "As perguntas de pesquisa e a literatura inicial forneceram dimensões sensibilizadoras, sem constituírem um "
    "esquema final fechado. A análise combinou codificação aberta, agrupamento axial, comparação constante e "
    "refinamento iterativo. A consolidação desses padrões resultou nas cinco camadas, posteriormente aplicadas "
    "como vocabulário normalizado ao corpus. As frequências descrevem a incidência dos mecanismos no esquema "
    "consolidado e não constituem validação empírica do modelo."
)

LLM_METHOD = (
    "A extração de texto completo registrou páginas, qualidade da extração e trechos relevantes. Uma triagem "
    "determinística localizou termos e evidências literais; em seguida, o LLM recebeu metadados e trechos "
    "selecionados, produziu campos estruturados e foi instruído a não inferir informação ausente. O modelo sugeriu "
    "elegibilidade, classificação e códigos, mas não tomou a decisão científica final."
)

LLM_VALIDATION = (
    "O autor confrontou as sugestões com o texto integral, verificou trechos e páginas e resolveu divergências "
    "pela leitura da evidência correspondente. A decisão fundamentada do autor prevaleceu. Não houve dupla revisão "
    "humana independente, e a concordância humano-LLM foi tratada apenas descritivamente. Permanecem riscos de "
    "enquadramento, interpretação, alucinação, falha de extração e dependência de versão, prompt e configuração; "
    "a confiança numérica do LLM não foi interpretada como probabilidade calibrada."
)

QUALITY_ROLE = (
    "Os instrumentos CASP/JBI adaptados foram usados como apoio para identificar limitações metodológicas e "
    "qualificar a interpretação, sem determinar elegibilidade ou a distinção entre evidência central e de apoio. "
    "As dimensões do CERQual orientaram a reflexão sobre coerência, adequação, relevância e limitações dos achados "
    "qualitativos; não foram atribuídos níveis formais de confiança a achados incompatíveis com essa abordagem."
)

LIMITATION = (
    "A generalização deve ser cautelosa diante da concentração setorial e da heterogeneidade de estudos empíricos, "
    "conceituais, normativos e técnicos. Mecanismos de ordenação, indexação e disponibilidade podem influenciar a "
    "recuperação; a combinação de múltiplas fontes, análise de posições posteriores, rastreamento bibliográfico e "
    "deduplicação reduz essa dependência, mas não demonstra exaustividade. A revisão por pesquisador único e a "
    "adjudicação assistida podem introduzir erro de classificação ou enquadramento. Validação literal, consulta ao "
    "texto integral e registro das decisões reduzem esses riscos, sem equivaler a dupla codificação independente."
)


def find(doc: Document, fragment: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if fragment in paragraph.text:
            return paragraph
    raise ValueError(f"Parágrafo não localizado: {fragment}")


def replace(doc: Document, fragment: str, text: str) -> Paragraph:
    paragraph = find(doc, fragment)
    run_properties = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs and paragraph.runs[0]._r.rPr is not None else None
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)
    return paragraph


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    element = OxmlElement("w:p")
    if paragraph._p.pPr is not None:
        element.append(deepcopy(paragraph._p.pPr))
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    run = created.add_run(text)
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run._r.insert(0, deepcopy(paragraph.runs[0]._r.rPr))
    return created


def replace_in_tables(doc: Document, fragment: str, text: str) -> None:
    replacements = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if fragment in paragraph.text:
                        if paragraph.runs:
                            paragraph.runs[0].text = text
                            for run in paragraph.runs[1:]:
                                run.text = ""
                        else:
                            paragraph.add_run(text)
                        replacements += 1
    if replacements == 0:
        raise ValueError(f"Texto de tabela não localizado: {fragment}")


def build(source: Path, target: Path) -> None:
    doc = Document(source)
    replace(doc, "Esta revisão sistemática investiga mecanismos", SUMMARY_PT)
    replace(doc, "This systematic review investigates conversational", SUMMARY_EN)
    replace(doc, "A busca combinou corpus-semente", SOURCES)
    replace_in_tables(doc, "Oito fontes, com sintaxe adaptada por API", "Cinco fontes, com sintaxe adaptada por serviço")
    replace(doc, "As consultas foram executadas em julho de 2026", RETRIEVAL)
    replace(doc, "O snowballing incluiu referências citadas", SELECTION)
    replace(doc, "Os critérios de elegibilidade selecionaram estudos", CORPUS)
    replace(doc, "Os 177 estudos estão identificados", TRACEABILITY)
    replace(doc, "A extração de texto completo registrou páginas", LLM_METHOD)
    replace(doc, "As evidências atribuídas pelo modelo foram confrontadas", LLM_VALIDATION)
    validation = find(doc, "O autor confrontou as sugestões")
    insert_after(validation, QUALITY_ROLE)
    synthesis = replace(doc, "A síntese temática identificou mecanismos", CENTRAL_CRITERIA)
    insert_after(synthesis, LAYER_ORIGIN)
    replace(doc, "A generalização deve ser cautelosa", LIMITATION)

    replace(doc, "Gráfico 1. Composição dos 407", "Gráfico 1. Composição das decisões documentadas na seleção integral")
    replace(
        doc,
        "Nota. As categorias representam o corpus analítico",
        "Nota. As categorias representam estudos incluídos, registros fora do escopo analítico e uma versão redundante vinculada à publicação final.",
    )
    replace(
        doc,
        "Compliance e gestão de risco apresentaram a maior cobertura",
        "Compliance e gestão de risco apresentaram a maior cobertura, com 339 estudos e 29 das 30 evidências centrais. "
        "Accountability e auditoria apareceram em 250 estudos, supervisão humana e escalonamento em 243, e aprendizagem "
        "operacional e monitoramento em 232. Contestabilidade e reparo tiveram incidência residual, com três estudos e "
        "nenhuma evidência central. O padrão desloca a lacuna principal da identificação de riscos para a capacidade "
        "de recurso, correção e reparação.",
    )
    replace(
        doc,
        "A camada técnica foi identificada em 147 estudos",
        "A camada regulatória foi identificada em 339 estudos e a organizacional em 324, seguidas pela técnica, com 284. "
        "As dimensões interacional e evolutiva reuniram 248 e 233 estudos. A diferença entre as camadas é menos polarizada "
        "que no corpus anterior, mas a raridade de mecanismos de contestação mostra que a presença de códigos interacionais "
        "não equivale, por si só, a poder efetivo de ação do usuário.",
    )
    replace(
        doc,
        "Mecanismos de supervisão humana e escalonamento foram identificados em 76 estudos",
        "Mecanismos de supervisão humana e escalonamento foram identificados em 243 estudos, incluindo 14 evidências "
        "centrais; contestabilidade e reparo apareceram em apenas três, sem evidência central. A diferença revela atenção "
        "ampla à intervenção humana, mas baixa operacionalização de mecanismos formais para questionar, revisar ou reparar respostas.",
    )
    replace(
        doc,
        "A classificação por domínio primário revelou concentração expressiva",
        "A classificação por domínio primário revelou concentração na saúde e na medicina, com 150 estudos (41,9% do "
        "corpus). Outros 112 (31,3%) apresentaram natureza multissetorial ou transversal, enquanto 47 (13,1%) se "
        "concentraram em tecnologia e operações empresariais.",
    )
    replace(
        doc,
        "Os demais ambientes regulados apresentaram cobertura consideravelmente menor",
        "Os demais ambientes regulados apresentaram cobertura menor: educação reuniu 17 estudos (4,7%); finanças e "
        "seguros, 14 (3,9%); infraestrutura crítica e cibersegurança, dez (2,8%); jurídico e judiciário, quatro (1,1%); "
        "e governo e setor público, quatro (1,1%).",
    )
    replace(
        doc,
        "A concentração é ainda mais acentuada entre as evidências centrais",
        "Entre as 30 evidências centrais, 19 pertencem à saúde e à medicina, seis a tecnologia e operações empresariais, "
        "quatro são multissetoriais e uma se concentra em finanças e seguros. Os demais domínios possuem estudos de apoio, "
        "mas nenhum classificado como evidência central.",
    )
    replace(
        doc,
        "A concentração em saúde e medicina, responsável por 44,1%",
        "A concentração em saúde e medicina, responsável por 41,9% do corpus, contrasta com a cobertura reduzida dos "
        "demais domínios específicos. Essa distribuição limita a transferência direta dos achados e exige validação do "
        "modelo em ambientes com obrigações, riscos, usuários, consequências e práticas profissionais distintos.",
    )
    replace(
        doc,
        "Observabilidade, auditoria e monitoramento pós-implantação foram identificados em 94 estudos",
        "Observabilidade, auditoria e monitoramento pós-implantação foram identificados em 276 estudos (77,1% do corpus), "
        "dos quais 23 centrais e 253 de apoio. Os estudos abordam auditoria, logs, tracing, telemetria, documentação, "
        "monitoramento contínuo e investigação de incidentes ao longo do ciclo de vida.",
    )
    replace(
        doc,
        "Governança do conhecimento, RAG e guardrails constituíram o achado de menor cobertura",
        "Conhecimento, RAG e guardrails foram associados a 160 estudos, incluindo 14 evidências centrais. A família mais "
        "estrita de governança do conhecimento apareceu em 44, mostrando que trabalhos sobre RAG ou guardrails nem sempre "
        "abordam curadoria, proveniência, validade, autoridade e versionamento das fontes.",
    )
    replace(
        doc,
        "Confiança, explicabilidade e governança orientada ao usuário foram identificadas em 101 estudos",
        "Confiança, explicabilidade e orientação ao usuário foram identificadas em 278 estudos (77,7% do corpus), incluindo "
        "17 evidências centrais e 261 de apoio. A ampla presença de transparência e comunicação de limites contrasta com "
        "a baixa incidência de mecanismos que permitem contestar ou reparar uma resposta.",
    )
    replace(
        doc,
        "A concentração da literatura em saúde e medicina, responsável por 44,1%",
        "A concentração da literatura em saúde e medicina, responsável por 41,9% do corpus e por 19 das 30 evidências "
        "centrais, limita a transferência direta dos achados. Finanças, educação, jurídico, governo e infraestrutura "
        "crítica têm menor presença de evidências centrais e exigem validação própria.",
    )
    replace(
        doc,
        "As maiores coocorrências ligam compliance e gestão de risco",
        "As maiores coocorrências ligam compliance e gestão de risco às camadas regulatória e organizacional; accountability "
        "e auditoria também se concentram nessas dimensões. A incidência residual de contestabilidade e reparo evidencia "
        "que cobertura institucional e técnica não implica, automaticamente, capacidade de ação do usuário.",
    )
    replace(
        doc,
        "A predominância das camadas técnica e organizacional reforça",
        "A predominância das camadas regulatória e organizacional reforça a natureza distribuída da accountability. Explicar "
        "e justificar condutas exige atores, instâncias de avaliação e consequências identificáveis (Bovens, 2007), enquanto "
        "a responsabilização algorítmica pode incidir sobre dados, modelos, decisões, efeitos e instituições (Wieringa, 2020). "
        "Para sistemas conversacionais, a cadeia relevante inclui ainda prompts, bases recuperadas, ferramentas, políticas e "
        "intervenção humana. O modelo proposto aproxima essa cadeia da auditoria de ciclo de vida defendida por Raji et al. "
        "(2020), conectando evidências técnicas a responsabilidades organizacionais.",
    )

    doc.core_properties.title = "Governança Conversacional em Sistemas Baseados em LLMs - v2.1"
    doc.core_properties.subject = "Correções metodológicas do terceiro parecer"
    doc.core_properties.comments = (
        "Versão v2.1. Recuperação ampliada, origem iterativa das camadas, critérios de evidência, adjudicação por LLM "
        "e corpus analítico reconciliados."
    )
    doc.save(target)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_article_v21.py SOURCE_V2.docx TARGET_V21.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
