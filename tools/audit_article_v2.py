from pathlib import Path
from docx import Document
import hashlib
import re
import sys
import zipfile


def content(doc):
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()] + [
        cell.text.strip() for table in doc.tables for row in table.rows for cell in row.cells
    ]


results = []
contents = []
for value in sys.argv[1:]:
    path = Path(value)
    doc = Document(path)
    items = content(doc)
    contents.append(items)
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    descriptions = re.findall(r'<wp:docPr[^>]*descr="([^"]*)"', xml)
    text = " ".join(items)
    results.append({
        "file": path.name,
        "paragraphs": sum(bool(p.text.strip()) for p in doc.paragraphs),
        "tables": len(doc.tables),
        "images": len(doc.inline_shapes),
        "alt_texts": sum(bool(x.strip()) for x in descriptions),
        "words": len(re.findall(r"\b\w+\b", text)),
        "text_sha256": hashlib.sha256(text.encode()).hexdigest(),
    })

for result in results:
    print(result)
if len(contents) == 2:
    print({"scientific_text_equal": contents[0] == contents[1], "items": [len(x) for x in contents]})
