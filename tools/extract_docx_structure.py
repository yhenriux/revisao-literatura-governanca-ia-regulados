"""Extrai a estrutura editorial de um DOCX sem modificar o arquivo-fonte."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def main() -> None:
    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    doc = Document(source)
    blocks = []
    for idx, block in enumerate(iter_blocks(doc), 1):
        if isinstance(block, Paragraph):
            blocks.append(
                {
                    "index": idx,
                    "type": "paragraph",
                    "style": block.style.name if block.style else "",
                    "text": block.text,
                }
            )
        else:
            blocks.append(
                {
                    "index": idx,
                    "type": "table",
                    "rows": [[cell.text for cell in row.cells] for row in block.rows],
                }
            )
    payload = {
        "source": str(source),
        "sections": len(doc.sections),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "blocks": blocks,
    }
    target.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
