from pathlib import Path
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TARGETS = [
    ROOT / "Artigo" / "Artigo_v2_final.docx",
    ROOT / "Artigo" / "Artigo_v2_para_editar.docx",
]


REPLACEMENTS = {
    "A busca combinou corpus-semente": (
        "A busca combinou corpus-semente, consultas automatizadas e expansão bibliográfica. Foram consultadas oito fontes: "
        "OpenAlex, Crossref, Semantic Scholar, PubMed, Europe PMC, CORE, arXiv e DOAJ. Cinco famílias cobriram governança "
        "de LLMs, LLMOps e observabilidade, governança conversacional, ambientes regulados e supervisão humana/contestabilidade. "
        "O suplemento registra, para cada uma das 40 combinações entre fonte e família, o limite solicitado, o estado de cobertura "
        "e os campos quantitativos preservados ou indisponíveis. A relação individual entre estudo, fonte e consulta não foi preservada "
        "e não foi reconstruída retrospectivamente."
    ),
    "As consultas foram executadas em julho de 2026": (
        "As consultas foram executadas em julho de 2026, sem filtros temporais ou linguísticos aplicados diretamente às APIs. O recorte "
        "principal de 2020 a 2026 foi aplicado na elegibilidade; trabalhos anteriores foram preservados apenas quando fundacionais. O limite "
        "de até 25 resultados por combinação foi uma restrição operacional para viabilizar registro, deduplicação, obtenção de texto completo "
        "e validação individual em oito fontes e cinco famílias. Ele não constitui amostragem probabilística, não neutraliza a ordenação das APIs "
        "e pode sub-representar consultas produtivas. Semantic Scholar e arXiv tiveram cobertura parcial por restrições de taxa e tempo de resposta. "
        "Como os totais retornados e armazenados por consulta não foram preservados, não é possível estimar retrospectivamente o impacto do limite; "
        "o snowballing reduziu, mas não eliminou, o risco de perda."
    ),
    "Os 177 estudos estão identificados individualmente": (
        "Os 177 estudos estão identificados individualmente por referência, classificação, PDF, hash e evidência literal no inventário suplementar. "
        "Na auditoria final, os 105 alertas históricos de localização foram reexaminados contra o texto integral: 94 âncoras foram confirmadas após "
        "normalização de Unicode, espaços e pontuação, e 11 foram substituídas por trechos alternativos literalmente localizados. Nenhum caso permaneceu "
        "sem evidência verificável. A reconciliação também documenta a duplicata removida e a readjudicação dos 17 casos de fronteira."
    ),
    "As evidências atribuídas pelo modelo foram confrontadas": (
        "As evidências atribuídas pelo modelo foram confrontadas com o texto extraído; correspondências ausentes e conflitos decisórios foram sinalizados. "
        "A auditoria final examinou todos os alertas remanescentes e preservou tanto o estado histórico quanto o trecho confirmado. Os 17 casos de fronteira "
        "receberam readjudicação humana documentada. Persistem riscos de interpretação, enquadramento, alucinação, falha de extração e dependência de versão, "
        "prompt e configuração; a confiança numérica do LLM não foi tratada como probabilidade calibrada. A qualidade foi avaliada por instrumento CASP/JBI "
        "adaptado, e a confiança analítica incorporou dimensões do CERQual (Aromataris et al., 2024; Lewin et al., 2018), sem exclusão automática por escore."
    ),
    "Accountability e auditoria foram identificadas": (
        "A incidência de accountability, auditoria, compliance e gestão de risco demonstra amplo reconhecimento de responsabilidades e obrigações, embora sua "
        "operacionalização varie entre documentação, controles técnicos, auditorias e estruturas organizacionais."
    ),
    "A síntese dos 177 estudos do corpus analítico resultou": None,
    "O primeiro achado foi sustentado por 118 estudos": (
        "A predominância do primeiro achado reflete a atenção dedicada a métricas, benchmarks, factualidade, alucinação, robustez, segurança e validação. "
        "Apesar dessa cobertura, a literatura permanece fragmentada na conversão dos riscos de modelos fundacionais em protocolos uniformes de avaliação "
        "e critérios de aceitação para sistemas conversacionais em produção (Bender et al., 2021; Bommasani et al., 2021; Weidinger et al., 2022)."
    ),
    "Supervisão humana e accountability operacional apareceram em 111 estudos": (
        "A diferença entre referências normativas à supervisão e mecanismos operacionais explícitos indica que accountability, oversight e controle humano são "
        "frequentemente defendidos sem detalhamento equivalente sobre atores, autoridade, evidências e responsabilidades."
    ),
    "Embora observabilidade, auditoria e monitoramento tenham aparecido em 94 estudos": (
        "A comparação com a camada evolutiva mostra que monitoramento é tratado sobretudo como rastreabilidade ou conformidade e menos como processo sistemático "
        "de aprendizagem, atualização e adaptação após incidentes."
    ),
    "A análise revela uma assimetria substantiva": (
        "A análise revela uma assimetria substantiva: a literatura privilegia mecanismos que informam o usuário sobre capacidades e limites, mas oferece menor "
        "cobertura para aqueles que permitem agir sobre uma resposta inadequada, obter revisão ou buscar reparação."
    ),
    "A generalização deve ser cautelosa": (
        "A generalização deve ser cautelosa porque a literatura se concentra em saúde e medicina e combina estudos empíricos, conceituais, normativos e técnicos. "
        "O limite por combinação, a cobertura parcial de duas fontes e a indisponibilidade dos totais históricos por consulta impedem demonstrar exaustividade ou "
        "quantificar o truncamento. A adjudicação assistida tampouco equivale a dupla revisão humana independente. Snowballing, validação literal dos 177 estudos e "
        "readjudicação dos casos de fronteira reduzem esses riscos, sem eliminá-los. Frequência temática, portanto, não prova implementação ou efetividade."
    ),
}


def replace_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def process(path):
    doc = Document(path)
    seen = set()
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        for prefix, replacement in REPLACEMENTS.items():
            if text.startswith(prefix):
                if prefix in seen:
                    raise RuntimeError(f"Prefixo duplicado em {path.name}: {prefix}")
                seen.add(prefix)
                if replacement is None:
                    delete_paragraph(paragraph)
                else:
                    replace_text(paragraph, replacement)
                break
    missing = set(REPLACEMENTS) - seen
    if missing:
        raise RuntimeError(f"Parágrafos não encontrados em {path.name}: {sorted(missing)}")
    doc.save(path)


for target in TARGETS:
    process(target)
    print(target)
