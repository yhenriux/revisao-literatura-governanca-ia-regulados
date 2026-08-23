#!/usr/bin/env python3
"""Executa a auditoria final reproduzível dos artefatos da v2.1."""

from __future__ import annotations

import csv
import hashlib
import json
import os
import re
import zipfile
from collections import Counter
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
ARTICLE = ROOT / "Artigo/Artigo_v2.1_para_editar.docx"
PDF = ROOT / "Artigo/Artigo_v2.1_para_leitura.pdf"
REDLINE = ROOT / "Artigo/Artigo_v2.1_com_alteracoes.docx"
BASE = ROOT / "Documentacao_do_projeto/v2.1/triagem_texto_completo"
CORPUS = BASE / "CORPUS_ANALITICO_FINAL_V2.1.csv"
DECISIONS = BASE / "REGISTRO_DECISOES_CORPUS_V2.1.csv"
MATRIX = BASE / "MATRIZ_ESTUDO_MECANISMO_CAMADA_V2.1.csv"
FIGURE_DATA = ROOT / "Recursos_do_artigo/v2.1/dados_figuras_v21.csv"
CATALOG = ROOT / "catalogo_virtual/catalogo.json"
OUTPUT = ROOT / "Documentacao_do_projeto/v2.1/RELATORIO_DE_QUALIDADE_V2.1.md"


def fs_path(path: Path) -> str:
    value = str(path.resolve())
    if os.name == "nt" and not value.startswith("\\\\?\\"):
        return "\\\\?\\" + value
    return value


def read_csv(path: Path) -> list[dict[str, str]]:
    with open(fs_path(path), encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def digest(path: Path) -> str:
    value = hashlib.sha256()
    with open(fs_path(path), "rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            value.update(block)
    return value.hexdigest()


def normalized_words(text: str) -> list[str]:
    return re.findall(r"[\wÀ-ÿ]+", text.casefold(), re.UNICODE)


def doc_text(doc: Document) -> str:
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    table_cells = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + table_cells)


def check(condition: bool, label: str, failures: list[str]) -> str:
    if not condition:
        failures.append(label)
        return "FALHOU"
    return "OK"


def main() -> int:
    failures: list[str] = []
    corpus = read_csv(CORPUS)
    decisions = read_csv(DECISIONS)
    matrix = read_csv(MATRIX)
    figure_data = read_csv(FIGURE_DATA)
    doc = Document(fs_path(ARTICLE))
    pdf = PdfReader(fs_path(PDF))
    text_doc = doc_text(doc)
    text_pdf = "\n".join(page.extract_text() or "" for page in pdf.pages)
    similarity = SequenceMatcher(None, normalized_words(text_doc), normalized_words(text_pdf), autojunk=False).ratio()

    classes = Counter(row["classificacao"] for row in corpus)
    missing_pdf = sum(not os.path.exists(fs_path(ROOT / row["arquivo_pdf"])) for row in corpus)
    missing_trace = sum(
        not row.get(field, "").strip()
        for row in corpus
        for field in ("id_estudo", "titulo", "autores", "ano", "arquivo_pdf", "hash_pdf", "pagina_evidencia", "trecho_evidencia")
    )
    graph1_total = sum(int(row["valor"]) for row in figure_data if row["figura"] == "grafico_1")
    with open(fs_path(CATALOG), encoding="utf-8") as handle:
        catalog = json.load(handle)

    with zipfile.ZipFile(fs_path(ARTICLE)) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    with zipfile.ZipFile(fs_path(REDLINE)) as archive:
        redline_xml = archive.read("word/document.xml").decode("utf-8")
    alt_descriptions = re.findall(r'<wp:docPr[^>]+descr="([^"]*)"', document_xml)
    tracked_insertions = redline_xml.count("<w:ins")
    tracked_deletions = redline_xml.count("<w:del")

    old_fragments = ["407 estudos", "177 estudos", "154 evidências", "25 primeiras"]
    results = [
        ("Corpus com 358 estudos", check(len(corpus) == 358, "corpus", failures)),
        ("30 evidências centrais", check(classes["evidencia_central"] == 30, "centrais", failures)),
        ("328 evidências de apoio", check(classes["evidencia_apoio"] == 328, "apoio", failures)),
        ("383 decisões documentadas", check(len(decisions) == 383, "decisões", failures)),
        ("PDFs localizados", check(missing_pdf == 0, "PDFs ausentes", failures)),
        ("Rastreabilidade completa", check(missing_trace == 0, "campos de rastreabilidade", failures)),
        ("Matriz longa existente", check(len(matrix) > len(corpus), "matriz longa", failures)),
        ("Gráfico 1 reconciliado", check(graph1_total == 383, "gráfico 1", failures)),
        ("Catálogo reconciliado", check(catalog.get("total_registros") == 358, "catálogo", failures)),
        ("Sete figuras no DOCX", check(len(doc.inline_shapes) == 7, "figuras", failures)),
        ("Três tabelas no DOCX", check(len(doc.tables) == 3, "tabelas", failures)),
        ("Textos alternativos", check(len(alt_descriptions) == 7 and all(alt_descriptions), "textos alternativos", failures)),
        ("PDF com 16 páginas", check(len(pdf.pages) == 16, "paginação", failures)),
        ("Equivalência DOCX/PDF", check(similarity >= 0.93, "equivalência DOCX/PDF", failures)),
        ("Redline rastreável", check(tracked_insertions > 0 and tracked_deletions > 0, "redline", failures)),
        ("Sem contagens substituídas", check(not any(fragment in text_doc for fragment in old_fragments), "contagens antigas", failures)),
    ]

    artifacts = [ARTICLE, PDF, REDLINE, CORPUS, DECISIONS, MATRIX, FIGURE_DATA, CATALOG]
    report = [
        "# Relatório de qualidade — artigo v2.1", "",
        "Data da auditoria: 23 de agosto de 2026.", "",
        "## Resultado", "",
        "| Critério | Estado |", "|---|---|",
    ]
    report.extend(f"| {label} | **{state}** |" for label, state in results)
    report += [
        "", "## Métricas", "",
        f"- Similaridade lexical entre DOCX e PDF: {similarity:.4f}.",
        f"- Palavras no DOCX, incluindo referências e tabelas: {len(normalized_words(text_doc))}.",
        f"- Linhas da matriz longa: {len(matrix)}.",
        f"- Alterações rastreadas: {tracked_insertions} inserções e {tracked_deletions} exclusões.",
        "", "## Hashes SHA-256", "",
        "| Artefato | SHA-256 |", "|---|---|",
    ]
    report.extend(f"| `{path.relative_to(ROOT).as_posix()}` | `{digest(path)}` |" for path in artifacts)
    report += ["", "## Conclusão", ""]
    if failures:
        report.append("A auditoria encontrou pendências: " + ", ".join(failures) + ".")
    else:
        report.append("Todos os critérios científicos, documentais e visuais verificáveis localmente foram atendidos. Resta apenas confirmar a sincronização remota após commit e publicação.")
    report.append("")
    with open(fs_path(OUTPUT), "w", encoding="utf-8", newline="\n") as handle:
        handle.write("\n".join(report))

    print(f"failures={len(failures)}")
    print(f"similarity={similarity:.4f}")
    print(f"pages={len(pdf.pages)} figures={len(doc.inline_shapes)} tables={len(doc.tables)}")
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
