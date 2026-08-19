"""Redesenha as figuras e a composição visual do artigo v2 sem alterar seu texto.

Uso:
    python tools/redesign_article_v2.py Artigo/Artigo_v2_final.docx
"""

from __future__ import annotations

import hashlib
import os
import re
import sys
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips


ROOT = Path(os.environ.get("REV_LIT_ROOT") or Path(__file__).resolve().parents[1])
PNG_DIR = ROOT / "Recursos_do_artigo" / "v2" / "imagens"
SVG_DIR = ROOT / "Recursos_do_artigo" / "v2" / "fontes_vetoriais"


FIGURES = [
    ("Grafico_1_composicao_do_corpus", 6.15, 1.80),
    ("Grafico_2_familias_de_mecanismos", 6.15, 3.45),
    ("Grafico_3_camadas_do_modelo", 6.15, 2.75),
    ("Grafico_4_distribuicao_setorial", 6.15, 3.25),
    ("Grafico_5_cobertura_dos_achados", 6.15, 2.75),
    ("Grafico_6_coocorrencia_mecanismos_camadas", 6.15, 3.55),
    ("Figura_1_modelo_de_cinco_camadas", 4.89, 3.57),
]

ALT_TEXTS = [
    "Barra segmentada com 407 registros: 177 estudos no corpus analítico, 112 referências contextuais e 118 estudos excluídos.",
    "Comparação entre incidência total e evidências centrais em oito famílias de mecanismos; compliance e gestão de risco têm a maior cobertura.",
    "Barras empilhadas com evidências centrais e de apoio nas cinco camadas; as camadas técnica e organizacional concentram mais estudos.",
    "Distribuição setorial do corpus analítico; saúde e medicina concentram 78 estudos e 12 evidências centrais.",
    "Comparação entre cobertura total e evidências centrais dos cinco achados da revisão.",
    "Mapa de calor da coocorrência entre oito famílias de mecanismos e cinco camadas de governança.",
    "Diagrama de governança integrada e retroalimentada com cinco camadas interdependentes ao redor de um sistema conversacional baseado em LLM.",
]

VISUAL_TEXT_REPLACEMENTS = {
    "Como os domínios são mutuamente exclusivos, a área ocupada por cada categoria representa sua participação no corpus analítico.":
        "Como os domínios são mutuamente exclusivos, o comprimento das barras representa sua participação no corpus analítico.",
    "Nota. A área de cada bloco é proporcional ao total de estudos. O número de evidências centrais é indicado em cada domínio.":
        "Nota. O comprimento das barras representa o total de estudos; o segmento em destaque indica as evidências centrais de cada domínio.",
    "Os cinco achados diferem tanto em cobertura temática quanto na quantidade de evidências centrais que os sustentam. O Gráfico 5 combina essas duas dimensões: a posição da bolha representa o total de estudos, e sua área indica o número de evidências centrais.":
        "Os cinco achados diferem tanto em cobertura temática quanto na quantidade de evidências centrais que os sustentam. O Gráfico 5 combina essas duas dimensões: o marcador final representa o total de estudos, e o marcador inicial indica o subconjunto de evidências centrais.",
    "Nota. Os achados são multirrótulo. A posição horizontal representa o total de estudos e a área da bolha representa a quantidade de evidências centrais.":
        "Nota. Os achados são multirrótulo. Os marcadores representam o total de estudos e o subconjunto de evidências centrais.",
}



def render_figures() -> None:
    """Valida os recursos previamente gerados pelo renderizador SVG/PNG."""
    for stem, width, height in FIGURES:
        png = PNG_DIR / f"{stem}.png"
        svg = SVG_DIR / f"{stem}.svg"
        requires_svg = stem.startswith("Grafico_")
        if not png.exists() or (requires_svg and not svg.exists()):
            raise FileNotFoundError(f"Recurso visual ausente: {stem}")
        with Image.open(png) as image:
            expected = (round(width * 300), round(height * 300))
            if image.width < expected[0] or image.height < expected[1]:
                raise ValueError(f"Resolução insuficiente em {png}: {image.size}; esperado ao menos {expected}")


def text_payload(doc: Document) -> str:
    paragraphs = "\n".join(p.text for p in doc.paragraphs)
    tables = "\n".join("|".join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows)
    return f"{paragraphs}\n{tables}"


def replace_visual_semantics(doc: Document) -> None:
    """Atualiza apenas frases que descreviam os formatos gráficos substituídos."""
    for paragraph in doc.paragraphs:
        for old, new in VISUAL_TEXT_REPLACEMENTS.items():
            if old not in paragraph.text:
                continue
            updated = paragraph.text.replace(old, new)
            if paragraph.runs:
                paragraph.runs[0].text = updated
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(updated)


def normalize_footer_page_numbers(doc: Document) -> None:
    """Grava um campo PAGE completo, incluindo separador e resultado em cache."""
    footers = []
    for section in doc.sections:
        # O documento usa rodapés distintos em páginas pares. Normalizar apenas
        # o rodapé padrão deixa o campo antigo ativo nas páginas 12 e 14.
        footers.extend((section.footer, section.even_page_footer))
    for footer in footers:
        paragraph = footer.paragraphs[0]
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.left_indent = Twips(0)
        paragraph.paragraph_format.right_indent = Twips(0)
        paragraph.paragraph_format.first_line_indent = Twips(0)

        begin_run = OxmlElement("w:r")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        begin_run.append(begin)
        paragraph._p.append(begin_run)

        instruction_run = OxmlElement("w:r")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        instruction_run.append(instruction)
        paragraph._p.append(instruction_run)

        separate_run = OxmlElement("w:r")
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        separate_run.append(separate)
        paragraph._p.append(separate_run)

        result_run = OxmlElement("w:r")
        value = OxmlElement("w:t")
        value.text = "1"
        result_run.append(value)
        paragraph._p.append(result_run)

        end_run = OxmlElement("w:r")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        end_run.append(end)
        paragraph._p.append(end_run)


def set_font(run, size: float = 8.5, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    rules = {
        "top": ("single", "8", "54707D"),
        "bottom": ("single", "8", "54707D"),
        "insideH": ("single", "4", "D5DEE3"),
        "left": ("nil", "0", "FFFFFF"),
        "right": ("nil", "0", "FFFFFF"),
        "insideV": ("nil", "0", "FFFFFF"),
    }
    for edge, (value, size, color) in rules.items():
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), value)
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    tbl_pr.append(borders)


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", total), ("tblInd", 0)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        cant_split = row._tr.get_or_add_trPr().find(qn("w:cantSplit"))
        if cant_split is None:
            row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Twips(width)


def format_tables(doc: Document) -> None:
    section = doc.sections[0]
    usable = int(section.page_width - section.left_margin - section.right_margin)
    ratios = ([0.07, 0.22, 0.49, 0.22], [0.10, 0.15, 0.75], [0.34, 0.23, 0.43])
    centers = ({0}, {0, 1}, {1})
    for table_index, (table, ratio, centered_columns) in enumerate(zip(doc.tables, ratios, centers)):
        widths = [round(usable * value) for value in ratio]
        widths[-1] += usable - sum(widths)
        set_table_geometry(table, widths)
        set_table_borders(table)
        header_props = table.rows[0]._tr.get_or_add_trPr()
        if header_props.find(qn("w:tblHeader")) is None:
            header_props.append(OxmlElement("w:tblHeader"))
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                if row_index == 0:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = tc_pr.find(qn("w:shd"))
                    if shd is None:
                        shd = OxmlElement("w:shd")
                        tc_pr.append(shd)
                    shd.set(qn("w:fill"), "E8F0F3")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if column_index in centered_columns else WD_ALIGN_PARAGRAPH.LEFT
                    fmt = paragraph.paragraph_format
                    fmt.left_indent = Twips(0)
                    fmt.right_indent = Twips(0)
                    fmt.first_line_indent = Twips(0)
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    fmt.line_spacing = 1.0
                    fmt.keep_together = True
                    for run in paragraph.runs:
                        set_font(run, size=8.5, bold=True if row_index == 0 else None, color="21343D")


def is_heading(paragraph) -> bool:
    text = paragraph.text.strip()
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    return style_name.startswith("heading") or bool(re.match(r"^\d+(?:\.\d+)*\.\s+", text)) or text in {"Resumo", "Abstract", "Status epistemológico do modelo", "Referências"}


def format_paragraphs(doc: Document) -> None:
    references = False
    previous_heading = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        fmt = paragraph.paragraph_format
        if re.fullmatch(r"(?:\d+\.\s*)?Referências", text):
            references = True
            # Convenção editorial acadêmica: a lista de referências inicia em
            # página própria. Além de melhorar a navegação, mantém a composição
            # final dentro do orçamento aprovado de 15 páginas.
            fmt.page_break_before = True
        caption = text.startswith(("Gráfico ", "Figura ", "Tabela "))
        note = text.startswith(("Nota.", "Fonte.", "Palavras-chave:", "Keywords:"))
        heading = is_heading(paragraph)
        fmt.left_indent = Twips(0)
        fmt.right_indent = Twips(0)
        if references and text and not heading:
            fmt.first_line_indent = Cm(-0.5)
            fmt.left_indent = Cm(0.5)
            fmt.line_spacing = 1.0
            # A bibliografia herdada estava em 8 pt. O ajuste para 10 pt melhora
            # a leitura de DOI, autoria e títulos sem inflar o corpo do artigo.
            for run in paragraph.runs:
                set_font(run, size=10)
        elif caption or note or heading or previous_heading or not text:
            fmt.first_line_indent = Twips(0)
        else:
            fmt.first_line_indent = Cm(0.5)
        if caption:
            fmt.keep_with_next = True
            fmt.space_after = Pt(3)
        if note:
            fmt.space_before = Pt(1)
        previous_heading = heading


def replace_figures(doc: Document) -> None:
    shapes = list(doc.inline_shapes)
    if len(shapes) != 7:
        raise ValueError(f"Esperadas 7 figuras; encontradas {len(shapes)}")
    for shape, (stem, width, height), alt in zip(shapes, FIGURES, ALT_TEXTS):
        image_path = PNG_DIR / f"{stem}.png"
        blip = shape._inline.graphic.graphicData.pic.blipFill.blip
        relation_id = blip.embed
        image_part = doc.part.related_parts[relation_id]
        image_part._blob = image_path.read_bytes()
        if stem == "Figura_1_modelo_de_cinco_camadas":
            for source_crop in shape._inline.graphic.graphicData.pic.blipFill.xpath("./a:srcRect"):
                source_crop.getparent().remove(source_crop)
        shape.width = Inches(width)
        shape.height = Inches(height)
        shape._inline.docPr.set("descr", alt)
        shape._inline.docPr.set("title", alt)
        parent = shape._inline
        while parent is not None and parent.tag != qn("w:p"):
            parent = parent.getparent()
        if parent is not None:
            from docx.text.paragraph import Paragraph

            paragraph = Paragraph(parent, doc)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Twips(0)
            paragraph.paragraph_format.right_indent = Twips(0)
            paragraph.paragraph_format.first_line_indent = Twips(0)
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_together = True
            # A legenda permanece junto da figura, mas a figura não força a nota,
            # a fonte e o parágrafo seguinte para a página seguinte. Isso evita
            # os vazios excessivos antes dos Gráficos 5 e 6.
            paragraph.paragraph_format.keep_with_next = False


def redesign(docx_path: Path) -> None:
    render_figures()
    doc = Document(docx_path)
    before_text = text_payload(doc)
    before_hash = hashlib.sha256(before_text.encode("utf-8")).hexdigest()
    expected_text = before_text
    for old, new in VISUAL_TEXT_REPLACEMENTS.items():
        expected_text = expected_text.replace(old, new)
    replace_visual_semantics(doc)
    normalize_footer_page_numbers(doc)
    format_tables(doc)
    format_paragraphs(doc)
    replace_figures(doc)
    after_text = text_payload(doc)
    after_hash = hashlib.sha256(after_text.encode("utf-8")).hexdigest()
    if after_text != expected_text:
        raise RuntimeError("O texto contém alterações além das quatro descrições visuais autorizadas.")
    temp_path = docx_path.with_name(f"{docx_path.stem}.visual_tmp.docx")
    marker = "Redesign visual v2: figuras em alta resolução e tabelas com geometria explícita; conteúdo científico preservado."
    comments = doc.core_properties.comments or ""
    if marker not in comments:
        doc.core_properties.comments = f"{comments} {marker}".strip()
    doc.save(temp_path)
    os.replace(temp_path, docx_path)
    print(f"text_sha256_before={before_hash}")
    print(f"text_sha256_after={after_hash}")
    print(f"docx={docx_path}")


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Uso: redesign_article_v2.py Artigo/Artigo_v2_final.docx")
    candidate = Path(sys.argv[1])
    if not candidate.is_absolute():
        candidate = Path.cwd() / candidate
    redesign(candidate)


if __name__ == "__main__":
    main()
