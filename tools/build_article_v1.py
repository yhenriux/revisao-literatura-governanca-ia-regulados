"""Constrói a v1 a partir da v0, preservando o pacote, estilos e figuras originais."""

from __future__ import annotations

import copy
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


INTRO = [
    "Modelos de Linguagem de Grande Escala (LLMs) ampliaram o uso de interfaces conversacionais em serviços de saúde, finanças, governo, jurídico, seguros, educação regulada e telecomunicações. A geração aberta de linguagem, a integração com fontes externas e a capacidade de acionar ferramentas aumentam o valor desses sistemas, mas também introduzem riscos de alucinação, opacidade, viés, uso indevido e instabilidade comportamental (Bender et al., 2021; Bommasani et al., 2021; Weidinger et al., 2022).",
    "Em ambientes regulados, uma resposta conversacional pode influenciar orientação clínica, aconselhamento financeiro, acesso a serviços públicos, interpretação jurídica ou exercício de direitos. A governança, portanto, não pode se restringir ao desempenho do modelo: precisa abranger a cadeia que produz a resposta, incluindo prompts, fontes de conhecimento, guardrails, ferramentas, registros, supervisão humana e contexto organizacional.",
    "A literatura oferece bases importantes, porém fragmentadas. Governança de IA e Responsible AI estabelecem princípios; accountability algorítmica trata de justificação, auditoria e responsabilização; estudos de LLMs examinam riscos e avaliação; interação humano-IA investiga confiança, transparência e correção; e trabalhos setoriais enfatizam conformidade e segurança. Falta uma síntese que integre essas perspectivas em mecanismos próprios da interação conversacional.",
    "Esta revisão sistemática identifica e organiza mecanismos técnicos, interacionais, organizacionais, regulatórios e evolutivos que orientam, controlam, monitoram, justificam e corrigem sistemas conversacionais baseados em LLMs em ambientes regulados. As questões de pesquisa examinam os mecanismos relatados, sua relação com risco e accountability, as capacidades associadas, as lacunas metodológicas e setoriais e o papel de explicabilidade, contestabilidade, reparo e aprendizagem operacional.",
    "A principal contribuição é o Modelo Conceitual Integrado de Governança Conversacional, apresentado desde o início como síntese dos resultados. Suas cinco camadas - técnica, interacional, organizacional, regulatória e evolutiva - mostram que governar a conversa requer coordenar controles do sistema, desenho da interação, responsabilidades institucionais, obrigações externas e aprendizagem em produção.",
    "O artigo contribui ao consolidar uma literatura dispersa, organizar os mecanismos em famílias analíticas e propor uma arquitetura conceitual aplicável à pesquisa, à avaliação e à prática organizacional. A narrativa segue do posicionamento da lacuna ao método, aos resultados, ao modelo e às suas implicações.",
]

RELATED = [
    "A literatura relevante converge em cinco vertentes. A primeira, governança de IA e Responsible AI, consolidou princípios como justiça, transparência, privacidade, segurança, robustez e accountability, mas também demonstrou que princípios isolados não garantem implementação responsável (Floridi et al., 2018; Jobin et al., 2019; Mittelstadt, 2019). Frameworks como o NIST AI RMF e o AI Act europeu aproximam esses princípios da gestão de risco, sem detalhar plenamente a governança que ocorre durante a conversa (National Institute of Standards and Technology, 2023; European Parliament & Council of the European Union, 2024).",
    "A segunda vertente trata de accountability, auditoria e explicabilidade. Accountability pressupõe atores capazes de explicar e justificar condutas diante de uma instância avaliadora, com possibilidade de julgamento e consequências (Bovens, 2007). Em sistemas de IA, o objeto dessa responsabilização se distribui entre dados, modelos, decisões, efeitos e instituições, exigindo documentação, rastreabilidade e auditoria ao longo do ciclo de vida (Raji et al., 2020; Wieringa, 2020).",
    "A terceira vertente examina modelos fundacionais e IA generativa. A escala e a generalidade dos LLMs propagam riscos por diferentes aplicações, enquanto respostas plausíveis podem ocultar erros factuais (Bommasani et al., 2021; Ji et al., 2023). RAG, guardrails, avaliação e observabilidade mitigam parte desses riscos, mas transferem novas responsabilidades para fontes, recuperação, integração e operação (Gao et al., 2023).",
    "A quarta vertente, interação humano-IA, mostra que usuários precisam compreender capacidades, limites, incertezas e caminhos de correção (Amershi et al., 2019; Shneiderman, 2020). Em uma interface conversacional, explicação, confirmação, recusa, escalonamento e reparo são simultaneamente escolhas de design e controles de governança; a fluência linguística pode elevar confiança além da competência real do sistema (Luger & Sellen, 2016; Rapp et al., 2021).",
    "A quinta vertente discute a adoção de IA em setores regulados. Saúde, finanças, governo, jurídico e seguros compartilham exigências de segurança, privacidade, auditabilidade, dever de cuidado e contestação, embora variem em risco e obrigação setorial. A lacuna não é ausência de princípios ou controles isolados, mas falta de integração entre o que controla o sistema, governa a interação, atribui responsabilidades, demonstra conformidade e aprende com o uso real. Essa lacuna fundamenta o modelo de cinco camadas.",
]

METHOD = [
    "A revisão seguiu o PRISMA 2020 e recomendações para revisões sistemáticas em engenharia de software, complementadas por snowballing (Kitchenham & Charters, 2007; Page et al., 2021; Wohlin, 2014). O protocolo articulou identificação e consolidação do corpus, avaliação de elegibilidade e qualidade, validação de evidências e síntese temática orientada à construção conceitual.",
    "A busca combinou corpus-semente, consultas automatizadas e expansão bibliográfica. Foram consultadas oito fontes: OpenAlex, Crossref, Semantic Scholar, PubMed, Europe PMC, CORE, arXiv e DOAJ. Cinco famílias de busca cobriram governança de LLMs, LLMOps e observabilidade, governança conversacional, ambientes regulados e supervisão humana/contestabilidade. As strings completas, adaptações por fonte, parâmetros de API, logs e hashes permanecem disponíveis no material suplementar e no repositório do projeto.",
    "As consultas foram executadas em julho de 2026, sem filtros temporais ou linguísticos aplicados diretamente às APIs. O recorte principal de 2020 a 2026 foi aplicado na elegibilidade; trabalhos anteriores foram preservados apenas quando fundacionais. A recuperação foi delimitada a até 25 resultados por combinação entre estratégia e fonte. Semantic Scholar e arXiv tiveram cobertura parcial por limitações de taxa e tempo de resposta.",
    "O snowballing incluiu referências citadas, trabalhos citantes, relacionados e expansões controladas por autoria e veículo. Todos os registros foram consolidados e deduplicados por DOI, título exato e similaridade textual, com validação de grupos potencialmente ambíguos. Após triagem e disponibilidade de texto completo, a busca foi congelada e 407 estudos únicos formaram o universo avaliado.",
    "Os critérios de elegibilidade selecionaram estudos sobre LLMs, IA generativa ou sistemas conversacionais que apresentassem mecanismos de governança e implicação para ambientes regulados ou de alto impacto. A avaliação distinguiu corpus analítico, referências fundacionais/contextuais e exclusões. O corpus analítico reuniu 177 estudos, dos quais 23 constituíram evidências centrais e 154 evidências de apoio; 112 referências foram mantidas para fundamentação e 118 estudos foram excluídos.",
    "A extração de texto completo registrou páginas, qualidade da extração e trechos relevantes, separando referências para reduzir contaminação bibliográfica. Uma triagem determinística em Python identificou termos, páginas e evidências literais. Em seguida, a adjudicação assistida por LLM classificou elegibilidade, qualidade e temas em JSON estruturado, limitada ao texto fornecido. O processo não correspondeu a revisão humana independente em duplicata; essa condição é tratada como limitação.",
    "Para reduzir extrapolações, evidências atribuídas pelo modelo foram verificadas contra o texto extraído. Casos sem correspondência ou em conflito com a decisão receberam atenção na adjudicação. A qualidade foi avaliada por instrumento CASP/JBI adaptado ao desenho dos estudos, e a confiança analítica incorporou dimensões do CERQual (Aromataris et al., 2024; Lewin et al., 2018). Os escores apoiaram priorização e cautela, sem operar como exclusão automática.",
    "A síntese temática identificou mecanismos, domínios e capacidades, normalizados por vocabulário controlado. A unidade de contagem foi o estudo deduplicado. Famílias, camadas e achados foram multirrótulo; por isso, suas frequências podem superar 177. O domínio primário foi mutuamente exclusivo. Somente o corpus analítico alimentou as frequências e o modelo conceitual.",
]

DISCUSSION = [
    "A revisão identificou uma literatura mais madura em compliance, risco, controles técnicos e avaliação do que em contestabilidade, reparo e aprendizagem operacional. Esse resultado converge com a crítica de que princípios de Responsible AI, embora amplamente consolidados, não garantem implementação responsável sem mecanismos institucionais e operacionais (Floridi et al., 2018; Jobin et al., 2019; Mittelstadt, 2019). No corpus analisado, a assimetria aparece justamente entre dimensões declaradas com frequência e capacidades menos operacionalizadas durante e após a interação.",
    "Os riscos atribuídos a modelos fundacionais - opacidade, alucinação, viés, conteúdo nocivo e dificuldade de avaliação - explicam a centralidade de guardrails, testes e observabilidade (Bender et al., 2021; Bommasani et al., 2021; Weidinger et al., 2022). A revisão também mostra, contudo, que controles do modelo não bastam. RAG transfere parte do risco para curadoria, proveniência e atualização das fontes; logs só geram governança quando alimentam auditoria; e monitoramento só produz aprendizagem quando incidentes resultam em mudança controlada (Gao et al., 2023).",
    "A predominância das camadas técnica e organizacional reforça a natureza distribuída da accountability. Explicar e justificar condutas exige atores, instâncias de avaliação e consequências identificáveis (Bovens, 2007), enquanto a responsabilização algorítmica pode incidir sobre dados, modelos, decisões, efeitos e instituições (Wieringa, 2020). Para sistemas conversacionais, a cadeia relevante inclui ainda prompts, bases recuperadas, ferramentas, políticas e intervenção humana. O modelo proposto aproxima essa cadeia da auditoria de ciclo de vida defendida por Raji et al. (2020), conectando evidências técnicas a responsabilidades organizacionais.",
    "A menor incidência das dimensões interacional e evolutiva evidencia uma lacuna que a literatura de interação humano-IA já permite antecipar: usuários precisam compreender capacidades, limites e possibilidades de correção (Amershi et al., 2019; Shneiderman, 2020). Fluência linguística pode elevar confiança sem elevar competência, especialmente quando o sistema não oferece escalonamento ou reparo (Luger & Sellen, 2016; Rapp et al., 2021). Assim, explicabilidade orientada ao usuário não se encerra na apresentação de uma justificativa; ela precisa permitir ação posterior, como contestar, corrigir, solicitar revisão ou acionar suporte humano.",
    "A revisão identificou mecanismos e padrões de incidência; o Modelo Conceitual Integrado, por sua vez, é uma proposição autoral derivada dessa síntese. Ele organiza relações entre as evidências, mas ainda não constitui uma escala de maturidade validada, uma demonstração de efetividade ou um substituto de normas setoriais. Sua utilidade está em tornar explícitas dependências que abordagens fragmentadas deixam dispersas e em oferecer categorias verificáveis para estudos empíricos, auditorias e desenho organizacional.",
    "A aplicação do modelo envolve trade-offs. Observabilidade amplia rastreabilidade, mas deve respeitar privacidade e minimização de dados; transparência pode favorecer confiança calibrada, mas não deve expor controles de segurança; supervisão humana reduz riscos apenas quando há autoridade, capacidade e tempo, podendo criar gargalos ou salvaguardas simbólicas; e padronização facilita auditoria, mas não elimina a necessidade de calibração por domínio. Esses conflitos impedem que uma camada seja maximizada isoladamente e reforçam a aplicação proporcional ao risco.",
    "Em saúde, interações clínicas exigem fontes validadas, supervisão profissional e escalonamento; em finanças, destacam-se rastreabilidade, prevenção de aconselhamento indevido e contestação; em governo, legitimidade, acesso a direitos e canais de recurso; e, em jurídico, seguros e educação regulada, a distinção entre informação, recomendação e decisão. O NIST AI RMF e o AI Act europeu sustentam uma abordagem baseada em risco, mas a tradução para a conversa depende do contexto e das obrigações de cada setor (National Institute of Standards and Technology, 2023; European Parliament & Council of the European Union, 2024).",
    "A generalização deve ser cautelosa porque a literatura está concentrada em saúde e medicina e combina estudos empíricos, conceituais, normativos e técnicos. A recuperação por API foi delimitada, algumas fontes tiveram cobertura parcial e o fluxo foi assistido por LLM sem revisores humanos independentes em duplicata. Validação literal e avaliação crítica reduzem, mas não eliminam, erros de extração, classificação ou perda de contexto. Frequência temática, portanto, não equivale a implementação nem a efetividade comprovada.",
    "Pesquisas futuras devem validar o modelo em organizações e setores distintos, comparar arranjos de supervisão, testar guardrails e RAG em produção e desenvolver métricas de contestabilidade, reparo e aprendizagem operacional. Estudos longitudinais podem examinar se logs, incidentes e feedback resultam em melhoria controlada. Essa validação deve combinar medidas de desempenho, análise de processos, experiência do usuário, incidentes e evidências de conformidade, permitindo distinguir presença formal de controles de sua efetividade observada. A agenda empírica deve avaliar não apenas se o sistema responde corretamente, mas se suas respostas podem ser reconstruídas, justificadas, auditadas, contestadas e corrigidas.",
]

CONCLUSION = [
    "A revisão demonstra que governança conversacional em ambientes regulados exige coordenação entre mecanismos técnicos, desenho da interação, responsabilidades organizacionais, obrigações regulatórias e aprendizagem operacional. Os achados respondem às questões de pesquisa ao identificar controles recorrentes e lacunas persistentes em contestabilidade, reparo, maturidade setorial e monitoramento orientado à melhoria.",
    "O Modelo Conceitual Integrado organiza essa evidência em cinco camadas interdependentes: técnica, interacional, organizacional, regulatória e evolutiva. Sua contribuição teórica é deslocar a unidade de análise do modelo isolado para o sistema conversacional sociotécnico; sua contribuição prática é oferecer uma estrutura proporcional ao risco para avaliar resposta, fontes de conhecimento, escalonamento, auditoria, contestação e mudança controlada.",
    "A interpretação deve considerar a concentração setorial da literatura, a heterogeneidade dos estudos e as limitações do fluxo assistido por LLM. O modelo ainda requer validação empírica e não substitui obrigações regulatórias específicas. Estudos futuros devem aplicá-lo em organizações de diferentes setores, comparar arranjos de supervisão e testar, em produção, se guardrails, RAG, observabilidade e canais de reparo produzem controle, accountability e aprendizagem verificáveis ao longo do ciclo de vida dos sistemas conversacionais em contextos reais de operação.",
]


RESULT_REPLACEMENTS = {
    146: "Nota. As categorias representam o corpus analítico, as referências fundacionais ou contextuais e as exclusões.",
    157: "As famílias de mecanismos e as camadas conceituais foram codificadas de maneira multirrótulo, de modo que um estudo pôde contribuir para diferentes categorias. As frequências representam incidência temática, e não necessariamente implementação ou validação empírica.",
    164: "Fonte. Elaboração própria com base no corpus analítico.",
    165: "Compliance e gestão de risco apresentaram a maior cobertura, com 157 estudos e todas as 23 evidências centrais. Controles técnicos e avaliação apareceram em 108 estudos, e accountability e auditoria, em 96. Contestabilidade e reparo tiveram incidência residual, com cinco estudos e uma evidência central. O padrão indica maior maturidade em prevenção, controle e conformidade do que em recurso, correção e reparação, deslocando a lacuna principal da identificação de riscos para a capacidade de resposta institucional.",
    170: "Fonte. Elaboração própria com base no corpus analítico.",
    171: "A camada técnica foi identificada em 147 estudos e a organizacional em 122, seguidas pela regulatória, com 85. As dimensões interacional e evolutiva tiveram menor incidência, com 49 e 36 estudos. A distribuição mostra maior consolidação da governança do sistema e da organização do que da governança do diálogo e da aprendizagem pós-implantação, duas dimensões essenciais para controlar o comportamento observado em uso.",
    183: "Mecanismos de supervisão humana e escalonamento foram identificados em 76 estudos, incluindo 12 evidências centrais; contestabilidade e reparo apareceram em apenas cinco, dos quais um central. A diferença revela maior atenção à presença de intervenção humana do que a mecanismos formais para questionar, revisar ou reparar respostas, indicando que supervisão e recurso ainda são tratados como capacidades distintas.",
    191: "Accountability e auditoria foram identificadas em 96 estudos, dos quais 17 centrais. Compliance e gestão de risco constituíram a família mais recorrente, alcançando todas as evidências centrais. A concentração mostra amplo reconhecimento de risco e conformidade, embora sua operacionalização varie entre documentação, auditoria, controles técnicos e estruturas organizacionais.",
    210: "Fonte. Elaboração própria com base no corpus analítico.",
    211: "A concentração em saúde e medicina, responsável por 44,1% do corpus, contrasta com a cobertura reduzida de finanças, educação, jurídico, infraestrutura crítica, telecomunicações e setor público. Essa distribuição limita a transferência direta dos achados e exige validação do modelo em ambientes regulados com obrigações, riscos, usuários, consequências e práticas profissionais distintos.",
    220: "Fonte. Elaboração própria com base no corpus analítico.",
    223: "O primeiro achado foi sustentado por 118 estudos, incluindo 18 evidências centrais, e teve a maior cobertura quantitativa. O resultado reflete a atenção dedicada a métricas, benchmarks, factualidade, alucinação, robustez, segurança e validação. Apesar dessa frequência, a literatura permanece fragmentada na conversão dos riscos de modelos fundacionais em protocolos uniformes de avaliação e critérios de aceitação para sistemas conversacionais em produção (Bender et al., 2021; Bommasani et al., 2021; Weidinger et al., 2022).",
    228: "Supervisão humana e accountability operacional apareceram em 111 estudos, mas mecanismos explícitos de supervisão ou escalonamento foram identificados em 76. A diferença indica que accountability, oversight e controle humano são frequentemente defendidos em nível normativo sem detalhamento equivalente sobre atores, autoridade, evidências e responsabilidades.",
    237: "Embora observabilidade, auditoria e monitoramento tenham aparecido em 94 estudos, a camada evolutiva foi identificada em apenas 36. Monitoramento é tratado sobretudo como rastreabilidade ou conformidade e menos como processo sistemático de aprendizagem, atualização e adaptação após incidentes.",
    240: "Governança do conhecimento, RAG e guardrails constituíram o achado de menor cobertura entre os cinco eixos, com 58 estudos. A família específica de governança do conhecimento apareceu em 44, mostrando que trabalhos sobre RAG ou guardrails nem sempre abordam curadoria, proveniência, validade, autoridade e versionamento das fontes.",
    248: "A análise revela uma assimetria substantiva: explicabilidade, confiança e comunicação de limites apareceram em 84 estudos, enquanto contestabilidade e reparo foram identificados em cinco. A literatura privilegia mecanismos que informam o usuário, mas oferece menor cobertura para aqueles que permitem agir sobre uma resposta inadequada, obter revisão ou buscar reparação.",
    252: "A concentração da literatura em saúde e medicina, responsável por 44,1% do corpus e por 12 das 23 evidências centrais, limita a transferência direta dos achados. Finanças, seguros, educação, jurídico, telecomunicações e infraestrutura crítica têm menor presença de evidências centrais e exigem validação própria.",
    259: "Fonte. Elaboração própria com base nos códigos normalizados do corpus analítico.",
    260: "As maiores coocorrências ligam compliance e gestão de risco às camadas técnica e organizacional; controles técnicos e avaliação também se concentram na camada técnica. Accountability e auditoria atravessam as dimensões técnica e organizacional, enquanto contestabilidade e reparo permanecem pouco representados. A matriz reforça a necessidade de integrar dimensões consolidadas a capacidades ainda pouco desenvolvidas, sem pressupor uma sequência linear entre as camadas.",
}

RESULT_OMIT = {156, 215, 221}


def paragraph_from_template(template, text: str, style: str | None = None):
    node = copy.deepcopy(template)
    texts = node.xpath('.//w:t')
    if texts:
        texts[0].text = text
        for item in texts[1:]:
            item.text = ''
    else:
        run = OxmlElement('w:r')
        txt = OxmlElement('w:t')
        txt.text = text
        run.append(txt)
        node.append(run)
    if style:
        ppr = node.find(qn('w:pPr'))
        if ppr is None:
            ppr = OxmlElement('w:pPr')
            node.insert(0, ppr)
        pstyle = ppr.find(qn('w:pStyle'))
        if pstyle is None:
            pstyle = OxmlElement('w:pStyle')
            ppr.insert(0, pstyle)
        pstyle.set(qn('w:val'), style)
    return node


def build(source: Path, target: Path) -> None:
    doc = Document(source)

    # Normalize regular and even-page footers. The v0 package carries an
    # even-page footer whose justified PAGE field drifts during PDF rendering.
    for section in doc.sections:
        for footer in (section.footer, section.even_page_footer):
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.clear()
            ppr = paragraph._p.get_or_add_pPr()
            jc = ppr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                ppr.append(jc)
            jc.set(qn('w:val'), 'center')
            paragraph.add_run('Página ')
            run = paragraph.add_run()
            begin = OxmlElement('w:fldChar')
            begin.set(qn('w:fldCharType'), 'begin')
            instruction = OxmlElement('w:instrText')
            instruction.set(qn('xml:space'), 'preserve')
            instruction.text = ' PAGE '
            end = OxmlElement('w:fldChar')
            end.set(qn('w:fldCharType'), 'end')
            run._r.extend((begin, instruction, end))
    body = doc.element.body
    original = [copy.deepcopy(el) for el in body.iterchildren() if el.tag != qn('w:sectPr')]
    sectpr = copy.deepcopy(body.sectPr)
    for el in list(body):
        body.remove(el)

    def original_block(number: int):
        return copy.deepcopy(original[number - 1])

    def add_text(text: str, kind: str = 'body'):
        template_number = {'title': 1, 'subtitle': 2, 'label': 3, 'body': 10, 'keywords': 5, 'h1': 9, 'h2': 23}[kind]
        style = {'h1': 'Heading1', 'h2': 'Heading2'}.get(kind)
        body.append(paragraph_from_template(original[template_number - 1], text, style))

    def add_original(numbers):
        for number in numbers:
            if number in RESULT_OMIT:
                continue
            if number in RESULT_REPLACEMENTS:
                body.append(paragraph_from_template(original[number - 1], RESULT_REPLACEMENTS[number]))
            else:
                body.append(original_block(number))

    add_text('Governança Conversacional em Sistemas Baseados em Modelos de Linguagem de Grande Escala em Ambientes Regulados: uma revisão sistemática da literatura', 'title')
    add_text('Conversational Governance in Large Language Model-Based Systems in Regulated Environments: A Systematic Literature Review', 'subtitle')
    add_text('Resumo', 'label')
    add_text('Esta revisão sistemática investiga mecanismos de governança conversacional em sistemas baseados em LLMs aplicados a ambientes regulados. Orientada pelo PRISMA 2020, avaliou 407 estudos em texto completo e consolidou 177 no corpus analítico. Os resultados mostram predominância de compliance, gestão de risco e controles técnicos, enquanto contestabilidade, reparo e aprendizagem operacional permanecem menos consolidados. A principal contribuição é um modelo integrado de cinco camadas - técnica, interacional, organizacional, regulatória e evolutiva - que desloca a governança do modelo isolado para o sistema conversacional sociotécnico.', 'body')
    add_text('Palavras-chave: governança conversacional; LLMs; ambientes regulados; accountability; auditoria; supervisão humana.', 'keywords')
    add_text('Abstract', 'label')
    add_text('This systematic review investigates conversational governance mechanisms in LLM-based systems deployed in regulated environments. Guided by PRISMA 2020, it assessed 407 full-text studies and consolidated 177 in the analytical corpus. Results show a predominance of compliance, risk management, and technical controls, whereas contestability, repair, and operational learning remain less consolidated. The main contribution is an integrated five-layer model - technical, interactional, organizational, regulatory, and evolutionary - that shifts governance from the isolated model to the sociotechnical conversational system.', 'body')
    add_text('Keywords: conversational governance; LLMs; regulated environments; accountability; auditing; human oversight.', 'keywords')

    add_text('1. Introdução', 'h1')
    for paragraph in INTRO:
        add_text(paragraph)

    add_text('2. Trabalhos relacionados', 'h1')
    for paragraph in RELATED:
        add_text(paragraph)

    add_text('3. Método', 'h1')
    for paragraph in METHOD[:2]:
        add_text(paragraph)
    add_original(range(69, 73))
    for paragraph in METHOD[2:5]:
        add_text(paragraph)
    add_original(range(109, 113))
    for paragraph in METHOD[5:]:
        add_text(paragraph)
    add_original(range(144, 148))

    add_text('4. Resultados e Modelo Conceitual Integrado', 'h1')
    add_original(range(156, 172))
    add_text('4.1. Mecanismos técnicos e operacionais', 'h2')
    add_original([173, 175, 176, 177, 178, 179, 180])
    add_text('4.2. Supervisão humana, escalonamento e contestabilidade', 'h2')
    add_original([183, 184, 186, 187, 188])
    add_text('4.3. Accountability, auditoria e compliance', 'h2')
    add_original([191, 192, 193, 194, 195, 196])
    add_text('4.4. Aplicações e domínios regulados', 'h2')
    add_original(range(203, 212))
    add_text('4.5. Achados da revisão', 'h2')
    add_original(range(213, 222))
    add_original([223, 224, 226, 228, 229, 231, 234, 235, 237, 240, 241, 243, 245, 246, 248, 249])
    add_text('4.6. Síntese dos achados', 'h2')
    add_original(range(251, 261))
    add_text('4.7. Modelo de cinco camadas', 'h2')
    add_original([268, 269, 276, 277, 278, 279, 280, 281])
    add_original(range(283, 291))
    add_original([292, 293, 295, 299, 300, 301, 302, 303, 304, 305, 306, 307])

    add_text('5. Discussão', 'h1')
    for paragraph in DISCUSSION:
        add_text(paragraph)

    add_text('6. Conclusão', 'h1')
    for paragraph in CONCLUSION:
        add_text(paragraph)

    add_text('Referências', 'h1')
    add_original(range(363, len(original) + 1))

    body.append(sectpr)
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)
    update_fields.set(qn('w:val'), 'true')
    doc.core_properties.title = 'Governança Conversacional em Sistemas Baseados em LLMs - v1'
    doc.core_properties.subject = 'Versão revisada após parecer editorial de Mauricio B. Almeida'
    doc.core_properties.comments = 'Derivada da v0; números e resultados preservados conforme decisão editorial.'
    doc.core_properties.modified = datetime.now(timezone.utc)

    figure_alts = [
        'Fluxo PRISMA com a composição dos 407 estudos avaliados em texto completo.',
        'Gráfico da incidência das famílias de mecanismos e da presença de evidência central.',
        'Gráfico da distribuição dos estudos pelas cinco camadas do modelo conceitual.',
        'Gráfico da composição setorial do corpus analítico.',
        'Gráfico da cobertura temática e da densidade de evidência central por achado.',
        'Mapa de calor da coocorrência entre famílias de mecanismos e camadas de governança.',
        'Modelo conceitual integrado de governança conversacional com cinco camadas interdependentes.',
    ]
    for shape, alt in zip(doc.inline_shapes, figure_alts):
        shape._inline.docPr.set('descr', alt)
        shape._inline.docPr.set('title', alt)
    doc.save(target)


if __name__ == '__main__':
    build(Path(sys.argv[1]), Path(sys.argv[2]))
