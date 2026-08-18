"""Gera a v2 final a partir da v1, com correções metodológicas e editoriais reproduzíveis."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.text.paragraph import Paragraph


RETRIEVAL = (
    "As consultas foram executadas em julho de 2026, sem filtros temporais ou linguísticos aplicados "
    "diretamente às APIs. O recorte principal de 2020 a 2026 foi aplicado na elegibilidade; trabalhos "
    "anteriores foram preservados apenas quando fundacionais. O limite de até 25 resultados por combinação "
    "entre estratégia e fonte foi adotado para tornar viáveis o registro, a deduplicação, a obtenção de texto "
    "completo e a validação individual em oito fontes e cinco famílias de busca. Esse limite não constitui "
    "amostragem probabilística e pode truncar consultas produtivas. Semantic Scholar e arXiv tiveram cobertura "
    "parcial por restrições de taxa e tempo de resposta; o snowballing reduziu, mas não eliminou, o risco de perda."
)

LLM_METHOD = (
    "A extração de texto completo registrou páginas, qualidade da extração e trechos relevantes, separando "
    "referências para reduzir contaminação bibliográfica. Uma triagem determinística em Python identificou "
    "termos, páginas e evidências literais. Em seguida, a adjudicação assistida por LLM recebeu metadados e "
    "trechos selecionados, produziu campos estruturados em JSON e foi instruída a não inferir informação "
    "ausente. O LLM funcionou como instrumento auxiliar, não como avaliador autônomo, e não houve revisão "
    "humana independente em duplicata."
)

LLM_VALIDATION = (
    "As evidências atribuídas pelo modelo foram confrontadas com o texto extraído; correspondências ausentes "
    "ou conflitos decisórios foram sinalizados, e os 17 casos de fronteira receberam readjudicação humana "
    "documentada. Persistem riscos de interpretação, enquadramento, alucinação, falha de extração e dependência "
    "de versão, prompt e configuração; a confiança numérica do LLM não foi tratada como probabilidade calibrada. "
    "A qualidade foi avaliada por instrumento CASP/JBI adaptado, e a confiança analítica incorporou dimensões do "
    "CERQual (Aromataris et al., 2024; Lewin et al., 2018), sem exclusão automática por escore."
)

STATUS = (
    "O Modelo Conceitual Integrado é uma proposição analítica derivada da síntese dos estudos. Ele não é uma "
    "escala de maturidade, norma, certificação ou requisito regulatório, nem substitui normas setoriais. A revisão "
    "identificou mecanismos e padrões de incidência; sua organização em cinco camadas é uma contribuição autoral. "
    "A validade externa e a efetividade do modelo dependem de validação empírica em organizações, sistemas reais "
    "e diferentes ambientes regulados."
)

INVENTORY_URL = (
    "https://github.com/yhenriux/revisao-literatura-governanca-ia-regulados/"
    "tree/main/Documentacao_do_projeto/methodology"
)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def find_paragraph(doc: Document, fragment: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if fragment in paragraph.text:
            return paragraph
    raise ValueError(f"Parágrafo não localizado: {fragment}")


def replace_paragraph(doc: Document, fragment: str, text: str) -> None:
    set_paragraph_text(find_paragraph(doc, fragment), text)


def remove_paragraph(doc: Document, fragment: str) -> None:
    paragraph = find_paragraph(doc, fragment)
    paragraph._element.getparent().remove(paragraph._element)


def insert_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    if text:
        created.add_run(text)
    return created


def copy_paragraph_format(target: Paragraph, source: Paragraph) -> None:
    """Replica a formatação compacta sem copiar o conteúdo do parágrafo-modelo."""
    source_ppr = source._p.pPr
    if source_ppr is not None:
        target_ppr = target._p.get_or_add_pPr()
        for child in list(target_ppr):
            target_ppr.remove(child)
        for child in source_ppr:
            target_ppr.append(deepcopy(child))


def copy_run_format(target_run, source_run) -> None:
    source_rpr = source_run._r.rPr
    if source_rpr is not None:
        target_run._r.insert(0, deepcopy(source_rpr))


def normalize_page_number_footer(doc: Document) -> None:
    """Mantém apenas o campo PAGE, evitando prefixo inconsistente na conversão do Drive."""
    for section in doc.sections:
        paragraph = section.footer.paragraphs[0]
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        run = OxmlElement("w:r")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run.extend((begin, instruction, end))
        paragraph._p.append(run)


def add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = (
        deepcopy(paragraph.runs[-1]._r.rPr)
        if paragraph.runs and paragraph.runs[-1]._r.rPr is not None
        else OxmlElement("w:rPr")
    )
    for tag in ("w:color", "w:u"):
        for existing in list(properties.findall(qn(tag))):
            properties.remove(existing)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def replace_table3(doc: Document) -> None:
    target = None
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == "Código":
            target = table
            break
    if target is None:
        raise ValueError("Tabela 3 não localizada")

    template = deepcopy(target.rows[1]._tr)
    for row in list(target.rows[1:]):
        target._tbl.remove(row._tr)

    rows = [
        ("I1", "Inclusão", "Objeto e mecanismo: LLM, IA generativa ou sistema conversacional com mecanismo de governança identificável."),
        ("I2", "Inclusão", "Contexto: aplicação em ambiente regulado ou de alto impacto, ou transferibilidade demonstrável para esse contexto."),
        ("I3", "Inclusão", "Evidência: texto completo suficiente, publicado entre 2020 e 2026; estudos anteriores somente quando fundacionais."),
        ("I4", "Inclusão", "Desenho aderente: estudo empírico, técnico, conceitual, normativo ou revisão com evidência substantiva."),
        ("E1", "Exclusão", "Escopo incompatível: ausência do sistema relevante ou de mecanismo de governança."),
        ("E2", "Exclusão", "Contexto insuficiente: sem aplicação, implicação ou transferibilidade para ambiente regulado ou de alto impacto."),
        ("E3", "Exclusão", "Evidência insuficiente: texto ou metadados inadequados para responder às questões da revisão."),
        ("E4", "Exclusão", "Redundância ou função inadequada: duplicata, versão redundante, manuscrito interno ou pré-2020 sem função fundacional."),
    ]
    for values in rows:
        target._tbl.append(deepcopy(template))
        row = target.rows[-1]
        for cell, value in zip(row.cells, values):
            set_paragraph_text(cell.paragraphs[0], value)
            for extra in list(cell.paragraphs[1:]):
                extra._element.getparent().remove(extra._element)


def add_traceability_statement(doc: Document) -> None:
    anchor = find_paragraph(doc, "O corpus analítico reuniu 177 estudos")
    paragraph = insert_after(anchor)
    copy_paragraph_format(paragraph, anchor)
    run = paragraph.add_run(
        "Os 177 estudos estão identificados individualmente por referência, classificação, PDF, hash e "
        "evidência-âncora no inventário suplementar. A reconciliação documenta a duplicata removida e a "
        "readjudicação dos 17 casos de fronteira. Os arquivos e o procedimento reproduzível estão no "
    )
    if anchor.runs:
        copy_run_format(run, anchor.runs[0])
    add_hyperlink(paragraph, "suplemento metodológico versionado", INVENTORY_URL)
    final_run = paragraph.add_run(".")
    if anchor.runs:
        copy_run_format(final_run, anchor.runs[0])


def add_epistemic_status(doc: Document) -> None:
    discussion = find_paragraph(doc, "5. Discussão")
    heading_template = find_paragraph(doc, "4.7. Modelo de cinco camadas")
    body_template = find_paragraph(doc, "Com base nos achados da revisão")
    heading = discussion.insert_paragraph_before("Status epistemológico do modelo")
    copy_paragraph_format(heading, heading_template)
    if heading.runs and heading_template.runs:
        copy_run_format(heading.runs[0], heading_template.runs[0])
    body = discussion.insert_paragraph_before(STATUS)
    copy_paragraph_format(body, body_template)
    if body.runs and body_template.runs:
        copy_run_format(body.runs[0], body_template.runs[0])


def build(source: Path, target: Path) -> None:
    doc = Document(source)

    normalize_page_number_footer(doc)

    replace_paragraph(
        doc,
        "A busca combinou corpus-semente",
        "A busca combinou corpus-semente, consultas automatizadas e expansão bibliográfica. Foram consultadas "
        "oito fontes: OpenAlex, Crossref, Semantic Scholar, PubMed, Europe PMC, CORE, arXiv e DOAJ. Cinco famílias "
        "cobriram governança de LLMs, LLMOps e observabilidade, governança conversacional, ambientes regulados e "
        "supervisão humana/contestabilidade. O suplemento preserva a estratégia em nível agregado, as limitações "
        "de cobertura e os artefatos de adjudicação; a relação individual entre estudo, fonte e consulta não foi "
        "preservada e não foi reconstruída retrospectivamente.",
    )
    replace_paragraph(
        doc,
        "Nota. As consultas foram adaptadas às regras de sintaxe",
        "Nota. As consultas foram adaptadas às regras de sintaxe, indexação e processamento de cada fonte; a "
        "cobertura operacional e suas limitações estão descritas no suplemento metodológico.",
    )
    replace_paragraph(doc, "As consultas foram executadas em julho de 2026", RETRIEVAL)
    replace_paragraph(doc, "A extração de texto completo registrou páginas", LLM_METHOD)
    replace_paragraph(doc, "Para reduzir extrapolações", LLM_VALIDATION)
    add_traceability_statement(doc)

    replace_table3(doc)
    replace_paragraph(
        doc,
        "Tabela 3. Critérios de inclusão e exclusão",
        "Tabela 3. Critérios consolidados de inclusão e exclusão",
    )
    replace_paragraph(
        doc,
        "Nota. O critério I7 permite",
        "Nota. Os critérios foram consolidados no corpo do artigo; as sete regras de inclusão, as sete de "
        "exclusão e os exemplos de fronteira permanecem integralmente no suplemento metodológico.",
    )

    replace_paragraph(
        doc,
        "A síntese quantitativa e temática demonstra",
        "A síntese integra os resultados em uma configuração desigual: controles que estabilizam modelos e "
        "responsabilidades institucionais estão mais consolidados que capacidades manifestadas no diálogo e "
        "na aprendizagem pós-incidente. Essa diferença orienta o modelo, mas não implica sequência linear ou "
        "escala de maturidade.",
    )
    remove_paragraph(doc, "Em conjunto, os dados indicam que a literatura não apresenta apenas")
    replace_paragraph(
        doc,
        "As capacidades identificadas não funcionam de forma isolada",
        "As capacidades identificadas não funcionam de forma isolada. RAG sem curadoria pode amplificar "
        "informação inadequada; guardrails sem monitoramento podem falhar silenciosamente; supervisão humana "
        "sem papéis definidos pode se tornar simbólica; logs sem processo de auditoria podem não gerar "
        "accountability; e explicabilidade sem um mecanismo de ação do usuário pode não produzir correção. A "
        "contribuição analítica da revisão está em demonstrar que a governança depende da integração das cinco "
        "dimensões, e não da presença isolada de controles.",
    )
    replace_paragraph(
        doc,
        "As maiores coocorrências ligam compliance",
        "As maiores coocorrências ligam compliance e gestão de risco às camadas técnica e organizacional; "
        "controles técnicos e avaliação também se concentram na camada técnica. A baixa coocorrência de "
        "mecanismos que permitem ao usuário agir evidencia uma dependência ainda pouco desenvolvida entre "
        "controle institucional e governança da interação, sem pressupor sequência linear entre as camadas.",
    )

    replace_paragraph(
        doc,
        "A revisão identificou uma literatura mais madura",
        "O desbalanceamento observado sugere que a institucionalização de Responsible AI ainda privilegia "
        "controles formalizáveis e responsabilidades internas, enquanto o poder de ação do usuário e a "
        "aprendizagem após falhas permanecem menos operacionalizados. Isso converge com a crítica de que "
        "princípios, embora consolidados, não garantem implementação responsável sem mecanismos institucionais "
        "e operacionais (Floridi et al., 2018; Jobin et al., 2019; Mittelstadt, 2019).",
    )
    replace_paragraph(
        doc,
        "A menor incidência das dimensões interacional e evolutiva",
        "A literatura de interação humano-IA ajuda a interpretar essa lacuna: usuários precisam compreender "
        "capacidades e limites e dispor de meios efetivos de correção (Amershi et al., 2019; Shneiderman, 2020). "
        "Como a fluência linguística pode elevar confiança sem elevar competência, explicabilidade orientada ao "
        "usuário deve permitir ação posterior - corrigir informações, solicitar revisão ou acionar suporte humano "
        "- e não apenas apresentar uma justificativa (Luger & Sellen, 2016; Rapp et al., 2021).",
    )
    replace_paragraph(
        doc,
        "A revisão identificou mecanismos e padrões de incidência; o Modelo Conceitual Integrado",
        "Essa distinção delimita a contribuição: os mecanismos e padrões de incidência derivam do corpus, "
        "enquanto o arranjo entre eles é uma síntese integradora. Sua utilidade está em tornar explícitas "
        "dependências dispersas e formular categorias verificáveis para estudos empíricos, auditorias e desenho "
        "organizacional, sem antecipar efetividade ainda não demonstrada.",
    )
    replace_paragraph(
        doc,
        "A generalização deve ser cautelosa",
        "A generalização deve ser cautelosa porque a literatura se concentra em saúde e medicina e combina "
        "estudos empíricos, conceituais, normativos e técnicos. A recuperação limitada por combinação, a "
        "cobertura parcial de duas fontes e a ausência de dupla revisão humana podem introduzir perda, erro de "
        "classificação ou dependência de enquadramento. Snowballing, validação literal e readjudicação dos casos "
        "de fronteira reduzem esses riscos, mas não demonstram exaustividade nem equivalem a dupla codificação. "
        "Frequência temática, portanto, não prova implementação ou efetividade.",
    )
    replace_paragraph(
        doc,
        "Pesquisas futuras devem validar o modelo",
        "Pesquisas futuras devem validar o modelo em organizações e setores distintos, comparar arranjos de "
        "supervisão, testar guardrails e RAG em produção e desenvolver métricas para ação do usuário e aprendizagem "
        "operacional. Estudos longitudinais podem examinar se logs, incidentes e feedback resultam em melhoria "
        "controlada. Essa validação deve combinar desempenho, análise de processos, experiência do usuário, "
        "incidentes e evidências de conformidade, distinguindo a presença formal de controles de sua efetividade.",
    )

    replace_paragraph(
        doc,
        "Nota. As perguntas e evidências são exemplificativas",
        "Nota. As perguntas sintetizam evidências recorrentes; sua organização em cinco camadas e seu uso como "
        "estrutura diagnóstica são proposições do modelo, ainda sujeitas a validação empírica. A aplicação deve "
        "ser calibrada conforme risco, setor e autonomia do sistema.",
    )

    replace_paragraph(
        doc,
        "A revisão demonstra que governança conversacional",
        "A revisão responde às questões de pesquisa ao mostrar que a governança conversacional depende da "
        "coordenação entre controles técnicos, desenho da interação, responsabilidades organizacionais, "
        "obrigações regulatórias e aprendizagem operacional. O problema central não é apenas controlar o modelo, "
        "mas governar o sistema sociotécnico que produz, apresenta, registra e corrige respostas.",
    )
    replace_paragraph(
        doc,
        "O Modelo Conceitual Integrado organiza essa evidência",
        "O Modelo Conceitual Integrado organiza essa evidência em cinco camadas interdependentes: técnica, "
        "interacional, organizacional, regulatória e evolutiva. Sua contribuição teórica é deslocar a unidade de "
        "análise do modelo isolado para o sistema conversacional; sua contribuição prática é oferecer categorias "
        "verificáveis para arquitetura, supervisão, auditoria, ação do usuário e mudança controlada.",
    )
    replace_paragraph(
        doc,
        "A interpretação deve considerar a concentração setorial",
        "A interpretação deve considerar a concentração setorial, a heterogeneidade dos estudos e as limitações "
        "da recuperação e da adjudicação assistida. O modelo é uma proposição ainda não validada empiricamente e "
        "não substitui obrigações setoriais. Estudos futuros devem aplicá-lo em organizações distintas, comparar "
        "arranjos de supervisão e testar se guardrails, RAG, observabilidade, revisão humana e aprendizagem após "
        "incidentes produzem controle e accountability verificáveis em operação.",
    )

    add_epistemic_status(doc)

    doc.core_properties.title = "Governança Conversacional em Sistemas Baseados em LLMs - v2 final"
    doc.core_properties.subject = "Versão metodológica final após auditoria de rastreabilidade"
    doc.core_properties.comments = (
        "Derivada da v1. Corpus reconciliado, Tabela 3 condensada, limitações reposicionadas e redundâncias reduzidas."
    )
    doc.save(target)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_article_v2.py SOURCE_V1.docx TARGET_V2.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
