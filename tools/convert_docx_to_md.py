"""Exporta o texto legível dos DOCX para Markdown antes da limpeza."""
from pathlib import Path
from docx import Document
import re
import sys

def clean(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()

def convert(src: Path, out_dir: Path) -> Path:
    doc = Document(src)
    lines = [f"# Exportação textual — {src.stem}", "", f"Origem: `{src.as_posix()}`", ""]
    for p in doc.paragraphs:
        text = clean(p.text)
        if not text:
            continue
        style = (p.style.name or "").lower()
        if "heading" in style or re.match(r"^\d+(?:\.\d+)*\.\s", text):
            lines += [f"## {text}", ""]
        else:
            lines += [text, ""]
    for i, table in enumerate(doc.tables, 1):
        lines += [f"## Tabela {i}", ""]
        rows = [[clean(c.text).replace("|", "\\|") for c in row.cells] for row in table.rows]
        if rows:
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
            lines.extend("| " + " | ".join(row) + " |" for row in rows[1:])
            lines.append("")
    out_dir.mkdir(parents=True, exist_ok=True)
    target = out_dir / f"{src.stem}.md"
    target.write_text("\n".join(lines), encoding="utf-8")
    return target

if __name__ == "__main__":
    root = Path.cwd()
    out = root / "Artigo" / "texto_exportado"
    for arg in sys.argv[1:]:
        print(convert(Path(arg), out))
