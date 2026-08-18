"""Formata somente o artefato final v2 conforme APA 7, sem alterar o conteúdo textual."""
from pathlib import Path
import re, sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BODY_HEADINGS = {
    '2. Trabalhos relacionados': 'Trabalhos relacionados',
    '3. Método': 'Método',
    '4. Resultados e Modelo Conceitual Integrado': 'Resultados e Modelo Conceitual Integrado',
    '5. Discussão': 'Discussão',
    '6. Conclusão': 'Conclusão',
    'Referências': 'Referências',
}

def set_cell_borders_minimal(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders'); tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = 'w:' + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); borders.append(el)
        el.set(qn('w:val'), 'single' if edge in ('top','bottom','insideH') else 'nil')
        el.set(qn('w:sz'), '4'); el.set(qn('w:space'), '0'); el.set(qn('w:color'), '808080')

def set_run_font(run, name='Times New Roman', size=12, bold=None, italic=None):
    run.font.name = name; run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def clear_direct_spacing(p):
    fmt = p.paragraph_format
    fmt.space_before = Pt(0); fmt.space_after = Pt(0); fmt.line_spacing = 2
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

def set_heading(p, level):
    p.style = f'Heading {level}'
    clear_direct_spacing(p)
    p.paragraph_format.first_line_indent = Inches(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs: set_run_font(r, bold=True, italic=(level == 3))

def build(source: Path, target: Path):
    doc = Document(source)
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = Inches(.5); sec.footer_distance = Inches(.5)

    styles = doc.styles
    normal = styles['Normal']; normal.font.name = 'Times New Roman'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); normal.font.size = Pt(12)
    for name, size in [('Heading 1',12),('Heading 2',12),('Heading 3',12)]:
        st=styles[name]; st.font.name='Times New Roman'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); st.font.size=Pt(size); st.font.bold=True
        st.paragraph_format.space_before=Pt(0); st.paragraph_format.space_after=Pt(0); st.paragraph_format.line_spacing=2

    paras = doc.paragraphs
    ref_start = None
    for idx, p in enumerate(paras):
        text = p.text.strip()
        if text in BODY_HEADINGS:
            p.text = BODY_HEADINGS[text]; set_heading(p, 1)
            if text == 'Referências': ref_start = idx
            continue
        m = re.match(r'^(\d+\.\d+)\.\s*(.+)$', text)
        if m:
            p.text = m.group(2); set_heading(p, 2); continue
        if re.match(r'^\d+\.\s+', text):
            # APA starts the introduction directly under the article title.
            if text.startswith('1. Introdução'):
                p._element.getparent().remove(p._element)
                continue
            p.text = re.sub(r'^\d+\.\s+', '', text); set_heading(p, 1); continue

        clear_direct_spacing(p)
        p.paragraph_format.first_line_indent = Inches(.5)
        if text in ('Resumo','Abstract'):
            p.paragraph_format.first_line_indent = Inches(0); set_heading(p, 1)
        elif text.startswith('Palavras-chave:') or text.startswith('Keywords:'):
            p.paragraph_format.first_line_indent = Inches(0)
            for r in p.runs: set_run_font(r, italic=False)
            if p.runs: p.runs[0].italic = True
        elif ref_start is not None and idx >= ref_start:
            p.paragraph_format.first_line_indent = Inches(-.5)
            p.paragraph_format.left_indent = Inches(.5)
        for r in p.runs: set_run_font(r)

    # Abstracts are single paragraphs without first-line indentation.
    refreshed = doc.paragraphs
    for p in refreshed:
        if p.text.strip().startswith('Esta revisão sistemática investiga') or p.text.strip().startswith('This systematic review investigates'):
            p.paragraph_format.first_line_indent = Inches(0)

    # Title block: centered, bold, readable, with no paragraph indent.
    for p in paras[:2]:
        clear_direct_spacing(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Inches(0)
        for r in p.runs: set_run_font(r, size=14 if p is paras[0] else 12, bold=(p is paras[0]))

    # Tables: APA-like horizontal rules, readable font, no paragraph indents.
    for table in doc.tables:
        set_cell_borders_minimal(table)
        for row_i, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    clear_direct_spacing(p); p.paragraph_format.first_line_indent = Inches(0)
                    for r in p.runs: set_run_font(r, size=10)
            if row_i == 0:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs: r.bold = True

    # Figures are centered; captions/notes remain left-aligned and double-spaced.
    for shape in doc.inline_shapes:
        shape._inline.docPr.set('descr', shape._inline.docPr.get('descr') or 'Figura do artigo; consultar a legenda para identificação.')
        shape._inline.docPr.set('title', shape._inline.docPr.get('title') or 'Figura do artigo')
    for p in doc.paragraphs:
        if re.match(r'^(Gráfico|Figura|Table|Figure)\s*\d+', p.text.strip()):
            p.paragraph_format.first_line_indent = Inches(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True

    for section in doc.sections:
        for footer in (section.footer, section.even_page_footer):
            if footer.paragraphs:
                footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.core_properties.title = 'Governança Conversacional em Sistemas Baseados em LLMs - v2 final APA 7'
    doc.core_properties.comments = 'Formatação visual APA 7 aplicada somente ao artefato final; conteúdo científico preservado.'
    doc.save(target)

if __name__ == '__main__':
    build(Path(sys.argv[1]), Path(sys.argv[2]))
